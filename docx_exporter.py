from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document

from table_extractor import TableData
from text_extractor import OCRPageResult


class DocxExportError(Exception):
    """Raised when DOCX export fails."""



def write_docx_results(
    text_results: Sequence[OCRPageResult],
    tables: Sequence[TableData],
    output_path: str | Path,
) -> Path:
    path = Path(output_path).expanduser().resolve()
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("OCR Result", level=1)

    tables_by_page: dict[int, list[TableData]] = {}
    for table in tables:
        tables_by_page.setdefault(table.page_number, []).append(table)

    for index, result in enumerate(text_results, start=1):
        document.add_heading(f"Page {result.page_number}", level=2)

        if result.text.strip():
            for paragraph in result.text.splitlines():
                document.add_paragraph(paragraph)
        else:
            document.add_paragraph("(No text detected)")

        page_tables = tables_by_page.get(result.page_number, [])
        for table in page_tables:
            document.add_paragraph(f"Table {table.table_index}")
            _append_table(document, table.rows)

        if index < len(text_results):
            document.add_page_break()

    document.save(path)
    return path



def _append_table(document: Document, rows: Sequence[Sequence[str]]) -> None:
    normalized_rows = list(rows) if rows else [[""]]
    column_count = max((len(row) for row in normalized_rows), default=1)
    table = document.add_table(rows=len(normalized_rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(normalized_rows):
        for col_index in range(column_count):
            value = row[col_index] if col_index < len(row) else ""
            table.cell(row_index, col_index).text = value
